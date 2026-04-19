import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


def _set_font(run, name="Times New Roman", size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def _add_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
                   size=12, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold)
    return p


def _add_separator(doc):
    """Thin horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


class DocxGenerator:
    def __init__(self, static_dir: str = "backend/static"):
        self.static_dir = static_dir
        os.makedirs(self.static_dir, exist_ok=True)

    def generate(self, doc_data: dict, extracted_facts: dict) -> tuple[str, bytes]:
        """
        Build a clean Word document from structured doc_data.
        Returns (filename, bytes).
        """
        try:
            doc = Document()

            # --- Page margins ---
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)

            # --- Header: WITHOUT THE LAAW BRANDING (keep professional) ---
            date_str = datetime.now().strftime("%d %B %Y")
            _add_paragraph(doc, date_str, alignment=WD_ALIGN_PARAGRAPH.RIGHT, size=11)

            doc.add_paragraph()  # blank line

            # --- Sender block ---
            sender = doc_data.get("sender_name", "").strip()
            sender_addr = doc_data.get("sender_address", "").strip()
            if sender:
                _add_paragraph(doc, sender, bold=True, space_after=0)
            if sender_addr:
                for line in sender_addr.split(","):
                    line = line.strip()
                    if line:
                        _add_paragraph(doc, line, size=11, space_after=0)

            doc.add_paragraph()

            # --- Recipient block ---
            recipient = doc_data.get("recipient_name", "").strip()
            recipient_addr = doc_data.get("recipient_address", "").strip()
            if recipient:
                _add_paragraph(doc, recipient, bold=True, space_after=0)
            if recipient_addr:
                for line in recipient_addr.split(","):
                    line = line.strip()
                    if line:
                        _add_paragraph(doc, line, size=11, space_after=0)

            doc.add_paragraph()
            _add_separator(doc)
            doc.add_paragraph()

            # --- Subject ---
            subject = doc_data.get("subject", "").strip().upper()
            if subject:
                _add_paragraph(doc, f"RE: {subject}", bold=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=12)

            _add_separator(doc)
            doc.add_paragraph()

            # --- Salutation ---
            salutation = doc_data.get("salutation", f"Dear {recipient or 'Sir/Madam'}").strip()
            _add_paragraph(doc, salutation, space_after=8)

            # --- Body paragraphs ---
            for para in doc_data.get("paragraphs", []):
                para = para.strip()
                if para:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(8)
                    run = p.add_run(para)
                    _set_font(run, size=12)

            # --- Demands / prayers ---
            demands = doc_data.get("demands", [])
            if demands:
                _add_paragraph(doc, "We therefore demand that you:", bold=True, space_before=4)
                for i, demand in enumerate(demands, 1):
                    demand = demand.strip()
                    if demand:
                        p = doc.add_paragraph(style="List Number")
                        p.paragraph_format.space_after = Pt(4)
                        run = p.add_run(demand)
                        _set_font(run, size=12)

            doc.add_paragraph()

            # --- Closing ---
            closing = doc_data.get("closing", "Yours faithfully").strip()
            _add_paragraph(doc, closing, space_after=36)

            # --- Signature ---
            _add_paragraph(doc, "_" * 30, space_after=2)
            if sender:
                _add_paragraph(doc, sender, bold=True, space_after=0)

            # --- Footer note ---
            doc.add_paragraph()
            _add_separator(doc)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("Prepared with TheLaaw • Nigerian Legal Aid Assistant")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

            # --- Save to bytes ---
            import io
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            docx_bytes = buf.read()

            doc_type = doc_data.get("document_type", "legal_document")
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            user_title = doc_data.get("subject", "").strip()
            if user_title:
                safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in user_title)[:50].strip()
                filename = f"{safe_title}_{timestamp}.docx"
            else:
                filename = f"{doc_type}_{timestamp}.docx"

            logger.info(f"Docx generated: {filename}")
            return filename, docx_bytes

        except Exception as e:
            logger.error(f"DocxGenerator error: {e}", exc_info=True)
            return None, None
